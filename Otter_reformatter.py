import streamlit as st
from docx import Document
from docx.shared import Pt
import re
import io

def reformat_transcript(input_file):
    """Reads a .docx transcript and reformats it."""
    doc = Document(input_file)
    new_doc = Document()

    # Set default font to Arial 10pt
    style = new_doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)

    current_speaker = None
    current_timestamp = None
    current_text = []

    # Pattern for lines like "Speaker 1  00:02"
    speaker_pattern = re.compile(r'^(Speaker\s*\d+)\s+(\d{2}:\d{2})')

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        match = speaker_pattern.match(text)
        if match:
            # Save previous speaker block before starting new one
            if current_speaker and current_text:
                combined = " ".join(current_text).strip()
                p = new_doc.add_paragraph()
                run_bold = p.add_run(f"[{current_timestamp}] {current_speaker.upper()}: ")
                run_bold.bold = True
                p.add_run(combined)
                current_text = []

            # Start a new block
            current_speaker = match.group(1)
            current_timestamp = match.group(2)
            remainder = text[match.end():].strip()
            if remainder:
                current_text.append(remainder)
        else:
            current_text.append(text)

    # Add the last block
    if current_speaker and current_text:
        combined = " ".join(current_text).strip()
        p = new_doc.add_paragraph()
        run_bold = p.add_run(f"[{current_timestamp}] {current_speaker.upper()}: ")
        run_bold.bold = True
        p.add_run(combined)

    output = io.BytesIO()
    new_doc.save(output)
    output.seek(0)
    return output

# Streamlit UI
st.set_page_config(page_title="Transcript Reformatter", page_icon="📝")
st.title("📝 Transcript Reformatter")
st.write("Upload a `.docx` transcript to clean and format it automatically. You can upload a new file at any time.")

uploaded_file = st.file_uploader("📂 Choose a .docx file", type=["docx"], key="uploader")

if uploaded_file:
    st.success(f"✅ File uploaded: {uploaded_file.name}")

    if st.button("Reformat Transcript"):
        with st.spinner("Processing... Please wait ⏳"):
            output_doc = reformat_transcript(uploaded_file)
        st.success("🎉 Done! Your transcript has been reformatted.")
        st.download_button(
            label="📥 Download Reformatted Transcript",
            data=output_doc,
            file_name="Reformatted_Transcript.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
# Write your code here :-)
